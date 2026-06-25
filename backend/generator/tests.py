import tempfile
import os
from io import BytesIO
from types import SimpleNamespace
from pathlib import Path
from unittest.mock import patch

from django.test import SimpleTestCase, override_settings
from django.utils import timezone
from docx import Document

from .docx_builder import _iter_all_paragraphs, build_generated_doc
from .views import _generation_download_name
from .keyword_matcher import score_resume
from .registry import requested_resume_dir, resume_registry_dir, save_bytes_to_resume_registry, save_file_to_resume_registry
from .services import (
    _build_refinement_prompt,
    _choose_better_openai_sections,
    _extract_contact,
    _normalize_sections,
    _skill_line,
    build_automatic_job_description,
    generate_resume_sections,
    render_resume_text,
)


class ResumeGenerationServiceTests(SimpleTestCase):
    def test_generation_download_name_includes_resume_job_position_and_time(self):
        created_at = timezone.datetime(2026, 5, 26, 21, 37, 19, tzinfo=timezone.get_current_timezone())
        generation = SimpleNamespace(
            created_at=created_at,
            resume=SimpleNamespace(title="Alex"),
            job_description=SimpleNamespace(company_name="Alan", job_title="Fullstack Software Engineer"),
        )

        self.assertEqual(
            _generation_download_name(generation, "docx"),
            "alex_alan_fullstack-software-engineer_20260526_213719.docx",
        )

    def test_automatic_job_description_uses_prompt_text_when_available(self):
        auto_job = build_automatic_job_description(
            "Senior Python engineer with Django and React experience.",
            "Target a backend-focused platform role with API design and cloud deployment.",
        )

        self.assertEqual(auto_job["job_title"], "Generated Role")
        self.assertIn("backend-focused platform role", auto_job["description_text"])

    def test_resume_registry_uses_resume_title_folder(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            generation = SimpleNamespace(id=7, resume=SimpleNamespace(title="Alex Chen Resume"))

            with override_settings(BASE_DIR=Path(temp_dir) / "backend"):
                target_dir = resume_registry_dir(generation)

            self.assertEqual(target_dir, Path(temp_dir) / "resume" / "alex-chen-resume")

    def test_resume_registry_saves_downloaded_files(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            generation = SimpleNamespace(id=7, resume=SimpleNamespace(title="Alex Chen Resume"))

            with override_settings(BASE_DIR=Path(temp_dir) / "backend"):
                pdf_path = save_bytes_to_resume_registry(generation, "resume.pdf", b"pdf-bytes")
                docx_path = save_file_to_resume_registry(generation, "resume.docx", BytesIO(b"docx-bytes"))

            self.assertEqual(pdf_path.read_bytes(), b"pdf-bytes")
            self.assertEqual(docx_path.read_bytes(), b"docx-bytes")

    def test_resume_registry_saves_to_requested_folder(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            generation = SimpleNamespace(id=7, resume=SimpleNamespace(title="Alex Chen Resume"))
            requested_folder = Path(temp_dir) / "requested"

            target_path = save_bytes_to_resume_registry(generation, "resume.pdf", b"pdf-bytes", requested_folder)

            self.assertEqual(target_path, requested_folder / "resume.pdf")
            self.assertEqual(target_path.read_bytes(), b"pdf-bytes")

    def test_requested_resume_dir_resolves_relative_to_project_root(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            with override_settings(BASE_DIR=Path(temp_dir) / "backend"):
                target_dir = requested_resume_dir("exports")

            self.assertEqual(target_dir, Path(temp_dir) / "exports")

    def test_extract_contact_ignores_plus_signs_in_resume_content(self):
        resume_text = "\n".join(
            [
                "Alex Chen",
                "alex@example.com Stockholm, Sweden",
                "SUMMARY",
                "Senior Engineer with 8+ years building SaaS platforms.",
                "SKILLS",
                "Frontend: React, JavaScript (ES6+), CSS3",
            ]
        )

        contact = _extract_contact(resume_text)

        self.assertIn("alex@example.com Stockholm, Sweden", contact)
        self.assertNotIn("8+ years", contact)
        self.assertNotIn("ES6+", contact)

    def test_build_generated_doc_does_not_write_summary_or_skills_as_contact(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            media_root = Path(temp_dir) / "media"
            template_path = Path(temp_dir) / "template.docx"
            document = Document()
            document.add_paragraph("Alex Chen")
            document.add_paragraph("Senior Full Stack Engineer")
            document.add_paragraph("alex@example.com Stockholm, Sweden")
            document.add_paragraph("")
            document.add_paragraph("SUMMARY")
            document.add_paragraph("Old summary")
            document.add_paragraph("SKILLS")
            document.add_paragraph("Old skills")
            document.add_paragraph("WORK EXPERIENCE")
            document.add_paragraph("Old experience")
            document.add_paragraph("EDUCATION")
            document.add_paragraph("Old education")
            document.save(template_path)

            sections = {
                "full_name": "Alex Chen",
                "contact_info": _extract_contact(
                    "\n".join(
                        [
                            "Alex Chen",
                            "alex@example.com Stockholm, Sweden",
                            "SUMMARY",
                            "Senior Engineer with 8+ years building SaaS platforms.",
                            "SKILLS",
                            "Frontend: React, JavaScript (ES6+), CSS3",
                        ]
                    )
                ),
                "professional_summary": "Generated summary",
                "technical_skills": "Generated skills",
                "professional_experience": "- Generated experience",
                "education": "Generated education",
            }

            with override_settings(MEDIA_ROOT=media_root):
                output_path = media_root / build_generated_doc(template_path, sections, 1)

            output = Document(output_path)
            texts = [p.text for p in _iter_all_paragraphs(output)]
            summary_index = texts.index("SUMMARY")
            header_text = "\n".join(texts[:summary_index])

            self.assertNotIn("8+ years", header_text)
            self.assertNotIn("ES6+", header_text)
            self.assertEqual(texts[summary_index + 1], "Generated summary")
            self.assertEqual(texts[texts.index("SKILLS") + 1], "- Generated skills")

    def test_local_generation_tailors_experience_to_job_and_prompt(self):
        resume_text = "\n".join(
            [
                "Alex Chen",
                "alex@example.com",
                "WORK EXPERIENCE",
                "Built React and TypeScript dashboards for financial analytics workflows.",
                "Designed Python APIs and PostgreSQL data services for reporting systems.",
                "Deployed Docker services on AWS with CI/CD pipelines and monitoring.",
            ]
        )
        job_description = "Need Python, PostgreSQL, AWS, Docker, CI/CD, analytics, dashboards, and API experience."
        custom_prompt = "Emphasize backend API delivery and cloud deployment."

        with patch.dict(os.environ, {"OPENAI_API_KEY": ""}):
            sections = generate_resume_sections(resume_text, job_description, custom_prompt)
        experience = sections["professional_experience"]
        generated_text = render_resume_text(sections)

        self.assertIn("Python", experience)
        self.assertIn("PostgreSQL", experience)
        self.assertIn("AWS", experience)
        self.assertIn("Docker", experience)
        self.assertNotEqual(experience, "Built React and TypeScript dashboards for financial analytics workflows.")
        self.assertGreaterEqual(score_resume(job_description, generated_text)["score"], 80)

    def test_local_generation_preserves_all_experience_companies(self):
        resume_text = "\n".join(
            [
                "Alex Chen",
                "alex@example.com",
                "WORK EXPERIENCE",
                "Finclude | Senior Full Stack Engineer | April 2022 - Present",
                "Built Python APIs and React dashboards for financial analytics workflows.",
                "MedIQ | Full Stack Engineer | September 2019 - March 2022",
                "Integrated healthcare platforms with Node.js and PostgreSQL services.",
                "Grid Finance | Full Stack Developer | June 2017 - August 2019",
                "Developed fintech dashboards and payment processing workflows.",
                "EDUCATION",
                "BSc Computer Science",
            ]
        )
        job_description = "Build AI-native internal tools with Python, React, PostgreSQL, workflows, and agents."

        with patch.dict(os.environ, {"OPENAI_API_KEY": ""}):
            sections = generate_resume_sections(resume_text, job_description, "")
        experience = sections["professional_experience"]

        self.assertIn("Finclude | Senior Full Stack Engineer | April 2022 - Present", experience)
        self.assertIn("MedIQ | Full Stack Engineer | September 2019 - March 2022", experience)
        self.assertIn("Grid Finance | Full Stack Developer | June 2017 - August 2019", experience)
        self.assertIn("supporting target-role priorities", experience)
        self.assertIn("Built Python APIs and React dashboards", experience)
        self.assertNotIn("BSc Computer Science", experience)

    def test_normalization_rebuilds_experience_when_ai_omits_companies(self):
        resume_text = "\n".join(
            [
                "Alex Chen",
                "alex@example.com",
                "WORK EXPERIENCE",
                "Finclude | Senior Full Stack Engineer | April 2022 - Present",
                "Built Python APIs and React dashboards.",
                "MedIQ | Full Stack Engineer | September 2019 - March 2022",
                "Integrated healthcare platforms with PostgreSQL services.",
                "EDUCATION",
                "BSc Computer Science",
            ]
        )
        ai_sections = {
            "full_name": "Alex Chen",
            "contact_info": "alex@example.com",
            "professional_summary": "Generated summary",
            "technical_skills": "Python, React, PostgreSQL",
            "professional_experience": "- Built AI-native internal tools with Python.",
            "projects": "- Project one\n- Project two",
            "education": "BSc Computer Science",
            "certifications": "",
        }

        normalized = _normalize_sections(
            ai_sections,
            resume_text,
            "Build AI-native tools with Python, React, and PostgreSQL.",
            "",
        )

        self.assertIn("Finclude | Senior Full Stack Engineer | April 2022 - Present", normalized["professional_experience"])
        self.assertIn("MedIQ | Full Stack Engineer | September 2019 - March 2022", normalized["professional_experience"])

    def test_technical_skills_are_grouped_resume_style(self):
        resume_text = """
        SKILLS
        Frontend: React.js, TypeScript, JavaScript (ES6+), React Native, HTML5, CSS3
        Backend: Node.js, Python, Express.js, REST APIs, GraphQL APIs
        Cloud & Infrastructure: Amazon Web Services (AWS), Docker, Kubernetes, CI/CD Pipelines
        Databases & Caching: PostgreSQL, Redis
        AI & Automation: AI-native systems, AI agents, retrieval workflows
        WORK EXPERIENCE
        Built React and Python systems on AWS.
        """
        job_description = "Need AI-native internal tools with Python, React, PostgreSQL, AWS, Docker, APIs, and agents."

        skills = _skill_line(resume_text, job_description, "")

        self.assertIn("- Frontend:", skills)
        self.assertIn("- Backend:", skills)
        self.assertIn("- Cloud & Infrastructure:", skills)
        self.assertIn("- Databases & Caching:", skills)
        self.assertIn("- AI & Automation:", skills)
        self.assertNotRegex(skills, r"^- [^:\n]+, [^:\n]+, [^:\n]+", "skills should not be one flat keyword bullet")

    def test_technical_skills_only_include_job_relevant_skills(self):
        resume_text = """
        SKILLS
        Frontend: React.js, TypeScript, JavaScript (ES6+), React Native, HTML5, CSS3
        Backend: Node.js, Python, Express.js, REST APIs, GraphQL APIs
        Cloud & Infrastructure: Amazon Web Services (AWS), Docker, Kubernetes, CI/CD Pipelines
        Databases & Caching: PostgreSQL, Redis
        Leadership: Technical Mentoring, Code Reviews
        """
        job_description = "Need backend Python APIs with PostgreSQL and Docker."

        skills = _skill_line(resume_text, job_description, "")

        self.assertIn("Python", skills)
        self.assertIn("PostgreSQL", skills)
        self.assertIn("Docker", skills)
        self.assertNotIn("React Native", skills)
        self.assertNotIn("HTML5", skills)
        self.assertNotIn("Technical Mentoring", skills)

    def test_broad_job_phrases_include_supported_related_skills(self):
        resume_text = """
        SKILLS
        Frontend: React.js, TypeScript, JavaScript (ES6+), HTML5, CSS3
        Backend: Python, REST APIs
        Databases & Caching: PostgreSQL
        AI & Automation: AI agents, automation, analytics dashboards
        """
        job_description = "Full-stack engineer using React, Python/Flask, PostgreSQL, LLM APIs, internal AI infrastructure, intelligent agents, and modern web technologies."

        skills = _skill_line(resume_text, job_description, "")

        self.assertIn("TypeScript", skills)
        self.assertIn("JavaScript (ES6+)", skills)
        self.assertIn("HTML5", skills)
        self.assertIn("CSS3", skills)
        self.assertIn("Python", skills)
        self.assertIn("PostgreSQL", skills)
        self.assertIn("AI Agents", skills)

    def test_refinement_prompt_includes_missing_keywords_and_source_headers(self):
        resume_text = "\n".join(
            [
                "WORK EXPERIENCE",
                "Finclude | Senior Full Stack Engineer | April 2022 - Present",
                "Built Python APIs.",
            ]
        )
        prompt = _build_refinement_prompt(
            {"professional_experience": "Finclude | Senior Full Stack Engineer | April 2022 - Present"},
            resume_text,
            "Need retrieval and LLMs.",
            "",
            {"score": 50, "missing_keywords": ["retrieval", "llms"]},
        )

        self.assertIn("retrieval", prompt)
        self.assertIn("llms", prompt)
        self.assertIn("Finclude | Senior Full Stack Engineer | April 2022 - Present", prompt)

    def test_refined_openai_sections_win_only_when_structure_is_preserved(self):
        resume_text = "\n".join(
            [
                "WORK EXPERIENCE",
                "Finclude | Senior Full Stack Engineer | April 2022 - Present",
                "Built Python APIs.",
            ]
        )
        job_description = "Need Python APIs, LLMs, and retrieval."
        draft = {
            "full_name": "Alex Chen",
            "contact_info": "alex@example.com",
            "professional_summary": "Python API engineer.",
            "technical_skills": "Python, APIs",
            "professional_experience": "Finclude | Senior Full Stack Engineer | April 2022 - Present\n- Built Python APIs.",
            "projects": "- API project\n- Backend project",
            "education": "BSc Computer Science",
            "certifications": "",
        }
        refined = {
            **draft,
            "professional_summary": "Python API engineer with LLM and retrieval experience.",
            "technical_skills": "Python, APIs, LLMs, Retrieval",
            "professional_experience": "Finclude | Senior Full Stack Engineer | April 2022 - Present\n- Built Python APIs for retrieval workflows and LLM integrations.",
        }

        chosen = _choose_better_openai_sections(draft, refined, resume_text, job_description, "")
        self.assertIn("LLM", chosen["professional_experience"])

        broken_refined = {**refined, "professional_experience": "- Built generic LLM systems."}
        chosen = _choose_better_openai_sections(draft, broken_refined, resume_text, job_description, "")
        self.assertIn("Finclude | Senior Full Stack Engineer | April 2022 - Present", chosen["professional_experience"])
