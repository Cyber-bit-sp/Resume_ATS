import re
from pathlib import Path

from django.conf import settings
from docx import Document
from docx.oxml.ns import qn as oxml_qn
from docx.text.paragraph import Paragraph as DocxParagraph


PLACEHOLDERS = {
    "{{FULL_NAME}}": "full_name",
    "{{CONTACT_INFO}}": "contact_info",
    "{{PROFESSIONAL_SUMMARY}}": "professional_summary",
    "{{TECHNICAL_SKILLS}}": "technical_skills",
    "{{PROFESSIONAL_EXPERIENCE}}": "professional_experience",
    "{{PROJECTS}}": "projects",
    "{{EDUCATION}}": "education",
    "{{CERTIFICATIONS}}": "certifications",
}

SECTION_LABELS = {
    "full_name": "Full Name",
    "contact_info": "Contact",
    "professional_summary": "Professional Summary",
    "technical_skills": "Technical Skills",
    "professional_experience": "Professional Experience",
    "projects": "Projects",
    "education": "Education",
    "certifications": "Certifications",
}

SECTION_HEADING_ALIASES = {
    "professional_summary": {"summary", "professional summary", "profile", "about", "about me", "objective", "career objective"},
    "technical_skills": {"skills", "technical skills", "core skills", "key skills", "competencies", "technologies", "tech stack"},
    "professional_experience": {"work experience", "professional experience", "experience", "employment history", "work history", "career history"},
    "projects": {"projects", "project experience", "key projects", "personal projects", "notable projects"},
    "education": {"education", "academic background", "academic history", "qualifications", "educational background"},
    "certifications": {"certifications", "certification", "licenses", "awards", "achievements", "credentials"},
}

EXPERIENCE_HEADER_PATTERN = re.compile(
    r"(\||\b(19|20)\d{2}\b|present|current|\d{1,2}/\d{4}|\d{4}\s*[-–]\s*(present|current|\d{4}))",
    re.IGNORECASE,
)


def _replace_in_paragraph(paragraph, values):
    original_text = paragraph.text
    updated_text = original_text
    for placeholder, key in PLACEHOLDERS.items():
        updated_text = updated_text.replace(placeholder, values.get(key, ""))
    if updated_text != original_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = updated_text
        else:
            paragraph.add_run(updated_text)
        return True
    return False


def build_generated_doc(template_path, sections, generation_id):
    template_path = Path(template_path) if template_path else None

    if not template_path or template_path.suffix.lower() != ".docx":
        document = Document()
        _rewrite_document_from_sections(document, sections)
        output_dir = settings.MEDIA_ROOT / "generated_resumes"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"final_resume_{generation_id}.docx"
        document.save(output_path)
        return output_path.relative_to(settings.MEDIA_ROOT)

    document = Document(template_path)
    replacements = 0

    for paragraph in document.paragraphs:
        replacements += int(_replace_in_paragraph(paragraph, sections))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replacements += int(_replace_in_paragraph(paragraph, sections))

    if replacements == 0:
        # Fallback for templates without placeholder tags: fill existing sections first.
        filled = _fill_existing_template_sections(document, sections)
        if not filled:
            _rewrite_document_from_sections(document, sections)

    output_dir = settings.MEDIA_ROOT / "generated_resumes"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"final_resume_{generation_id}.docx"
    document.save(output_path)
    return output_path.relative_to(settings.MEDIA_ROOT)


def _rewrite_document_from_sections(document, sections):
    title_style, heading_style, body_style, bullet_style = _detect_styles(document)

    body = document.element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body.remove(child)

    full_name = (sections.get("full_name") or "").strip()
    if full_name:
        _add_paragraph(document, full_name, title_style)

    for key in [
        "contact_info",
        "professional_summary",
        "technical_skills",
        "professional_experience",
        "projects",
        "education",
        "certifications",
    ]:
        value = (sections.get(key) or "").strip()
        if not value:
            continue
        _add_paragraph(document, SECTION_LABELS[key], heading_style)
        for line in value.splitlines():
            text = line.strip()
            if text:
                if text.startswith("-"):
                    _add_paragraph(document, text.lstrip("- "), bullet_style)
                else:
                    _add_paragraph(document, text, body_style)


def _add_paragraph(document, text, style_name=None):
    paragraph = document.add_paragraph(text)
    if style_name:
        try:
            paragraph.style = style_name
        except Exception:
            pass
    return paragraph


def _iter_all_paragraphs(document):
    """Return all paragraphs in document order, including those inside table cells."""
    result = []
    body = document.element.body
    W_P = oxml_qn('w:p')
    W_TBL = oxml_qn('w:tbl')
    W_TR = oxml_qn('w:tr')
    W_TC = oxml_qn('w:tc')

    for element in body:
        if element.tag == W_P:
            result.append(DocxParagraph(element, document))
        elif element.tag == W_TBL:
            for tr in element.iter(W_TR):
                for tc in tr:
                    if tc.tag == W_TC:
                        for p in tc:
                            if p.tag == W_P:
                                result.append(DocxParagraph(p, document))
    return result


def _detect_styles(document):
    paragraphs = [p for p in _iter_all_paragraphs(document) if (p.text or "").strip()]
    if not paragraphs:
        return None, None, None, None

    title_style = paragraphs[0].style.name if paragraphs[0].style else None
    heading_style = None
    body_style = None
    bullet_style = None

    for p in paragraphs:
        text = (p.text or "").strip()
        style_name = p.style.name if p.style else "Normal"
        style_lower = style_name.lower()
        # Prefer explicit Word heading styles
        if not heading_style and "heading" in style_lower:
            heading_style = style_name
        if not bullet_style and ("list" in style_lower or _paragraph_has_list_format(p)):
            bullet_style = style_name
        # Content heuristics as fallback
        if not heading_style and text and len(text) <= 40 and (text.isupper() or text.endswith(":")):
            heading_style = style_name
        if not body_style and len(text) > 40:
            body_style = style_name
        if not bullet_style and (text.startswith("-") or text.startswith("•")):
            bullet_style = style_name

    return title_style, heading_style or body_style, body_style or "Normal", bullet_style or "List Paragraph"


def _set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _normalize_heading_text(value):
    return " ".join((value or "").strip().lower().replace(":", "").split())


def _paragraph_style_name(paragraph):
    return paragraph.style.name if paragraph.style else None


def _paragraph_has_list_format(paragraph):
    try:
        ppr = paragraph._p.pPr
        return bool(ppr is not None and ppr.numPr is not None)
    except Exception:
        return False


def _clear_list_format(paragraph):
    try:
        ppr = paragraph._p.get_or_add_pPr()
        if ppr is not None and ppr.numPr is not None:
            ppr.remove(ppr.numPr)
    except Exception:
        return


def _remove_paragraph(paragraph):
    try:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception:
        return


def _find_heading_indices(paragraphs):
    indices = {}
    for index, paragraph in enumerate(paragraphs):
        text = _normalize_heading_text(paragraph.text)
        if not text:
            continue
        for section_key, aliases in SECTION_HEADING_ALIASES.items():
            if section_key in indices:
                continue
            if text in aliases:
                indices[section_key] = index
    return indices


def _append_p_after(ref_p, text, style_p=None, copy_list_from_p=None):
    """Insert a new paragraph immediately after ref_p in its XML parent. Returns the new raw p element."""
    from docx.oxml import OxmlElement
    from lxml import etree
    new_p = OxmlElement('w:p')
    # Copy paragraph properties from style_p if provided
    if style_p is not None and style_p.pPr is not None:
        new_p.append(etree.fromstring(etree.tostring(style_p.pPr)))
    if copy_list_from_p is not None:
        try:
            src_ppr = copy_list_from_p.pPr
            if src_ppr is not None and src_ppr.numPr is not None:
                tgt_ppr = new_p.get_or_add_pPr()
                if tgt_ppr.numPr is not None:
                    tgt_ppr.remove(tgt_ppr.numPr)
                tgt_ppr.append(etree.fromstring(etree.tostring(src_ppr.numPr)))
        except Exception:
            pass
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_p.append(r)
    ref_p.addnext(new_p)
    return new_p


def _write_section_between_headings(paragraphs, start_idx, end_idx, lines, section_key):
    editable_paragraphs = [paragraphs[idx] for idx in range(start_idx + 1, end_idx)]
    if not editable_paragraphs and not lines:
        return

    content_style = _paragraph_style_name(editable_paragraphs[0]) if editable_paragraphs else None

    # Find a template bullet paragraph to use as a source for list formatting
    bullet_source_paragraph = next(
        (p for p in editable_paragraphs if _paragraph_has_list_format(p)), None
    )

    for i, paragraph in enumerate(editable_paragraphs):
        if i < len(lines):
            line_text = lines[i]
            cleaned = line_text.lstrip("-• ").strip()
            is_experience_header = section_key == "professional_experience" and bool(EXPERIENCE_HEADER_PATTERN.search(cleaned))
            target_has_list = _paragraph_has_list_format(paragraph)
            if is_experience_header and target_has_list:
                _clear_list_format(paragraph)
                line_text = cleaned
            elif line_text.startswith("-"):
                if target_has_list:
                    # Preserve native bullet formatting — don't add a literal dash
                    line_text = cleaned
                else:
                    line_text = f"- {cleaned}"
            _set_paragraph_text(paragraph, line_text)
        else:
            _remove_paragraph(paragraph)

    if len(lines) > len(editable_paragraphs):
        # Always insert after the LAST paragraph already in this section (same XML parent/cell).
        # Never insert before the next section heading — it may live in a different table cell,
        # which would put the overflow content in the wrong section.
        if editable_paragraphs:
            ref_p = editable_paragraphs[-1]._p
        else:
            # No editable paragraphs at all — insert right after the heading itself.
            ref_p = paragraphs[start_idx]._p

        style_p = editable_paragraphs[-1]._p if editable_paragraphs else None
        for extra_line in lines[len(editable_paragraphs):]:
            cleaned = extra_line.lstrip("-• ").strip()
            is_experience_header = section_key == "professional_experience" and bool(EXPERIENCE_HEADER_PATTERN.search(cleaned))
            is_bullet_line = extra_line.startswith("-")
            bsrc = bullet_source_paragraph._p if (is_bullet_line and not is_experience_header and bullet_source_paragraph) else None
            ref_p = _append_p_after(ref_p, cleaned, style_p=style_p, copy_list_from_p=bsrc)


def _fill_header_area(paragraphs, first_heading_idx, sections):
    full_name = (sections.get("full_name") or "").strip()
    contact_raw = (sections.get("contact_info") or "").strip()

    if not full_name and not contact_raw:
        return

    # Include ALL paragraphs before the first heading (even empty ones — they are
    # intentional placeholder slots in the template header area).
    header_slots = list(paragraphs[:first_heading_idx])

    # If the template has NO paragraphs before the first heading, insert them now.
    if not header_slots:
        from docx.oxml import OxmlElement
        first_heading_p = paragraphs[first_heading_idx]._p
        # Build paragraphs and insert BEFORE the first heading (addprevious keeps order).
        name_p = OxmlElement('w:p')
        if full_name:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = full_name
            r.append(t)
            name_p.append(r)
        first_heading_p.addprevious(name_p)
        if contact_raw:
            contact_p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = contact_raw
            r.append(t)
            contact_p.append(r)
            name_p.addnext(contact_p)
        return

    # --- We have existing header slots ---
    # Split contact_info into pipe-separated parts
    contact_parts = [p.strip() for p in re.split(r"\s*\|\s*", contact_raw) if p.strip()]

    # Slot 0 → full name
    if full_name:
        _set_paragraph_text(header_slots[0], full_name)

    # Remaining slots → contact info parts
    remaining_slots = header_slots[1:]
    if not remaining_slots or not contact_parts:
        return

    if len(remaining_slots) == 1:
        # One slot — join everything
        _set_paragraph_text(remaining_slots[0], " | ".join(contact_parts))
    elif len(remaining_slots) >= len(contact_parts):
        # Enough slots for one part each
        for i, part in enumerate(contact_parts):
            _set_paragraph_text(remaining_slots[i], part)
    else:
        # More parts than slots — split: put non-email parts in slot 0, email/link parts in slot 1
        email_parts = [p for p in contact_parts
                       if "@" in p or "linkedin" in p.lower() or "github" in p.lower() or "http" in p.lower()]
        other_parts = [p for p in contact_parts if p not in email_parts]
        _set_paragraph_text(remaining_slots[0],
                            " | ".join(other_parts) if other_parts else " | ".join(contact_parts))
        if len(remaining_slots) >= 2:
            _set_paragraph_text(remaining_slots[1],
                                " | ".join(email_parts) if email_parts else "")


def _all_heading_texts():
    texts = set()
    for aliases in SECTION_HEADING_ALIASES.values():
        texts.update(aliases)
    return texts


def _section_lines(section_key, raw_value):
    value = (raw_value or "").strip()
    if not value:
        return []
    heading_texts = _all_heading_texts()
    lines = []
    for line in value.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        # Stop processing if we hit another section heading (AI mixed sections together)
        if _normalize_heading_text(stripped.lstrip("-• ")) in heading_texts:
            break
        lines.append(stripped)
    if section_key == "professional_experience":
        output = []
        for line in lines:
            cleaned = line.lstrip("-• ").strip()
            if EXPERIENCE_HEADER_PATTERN.search(cleaned):
                output.append(cleaned)
            else:
                output.append(line if line.startswith("-") else f"- {cleaned}")
        return output
    if section_key == "projects":
        return [line if line.startswith("-") else f"- {line}" for line in lines]
    if section_key == "technical_skills":
        return [line if line.startswith("-") else f"- {line}" for line in lines]
    return lines


def _fill_existing_template_sections(document, sections):
    # Use all paragraphs including those inside table cells (handles table-based templates)
    paragraphs = _iter_all_paragraphs(document)
    if not paragraphs:
        return False

    heading_indices = _find_heading_indices(paragraphs)
    if len(heading_indices) < 2:
        return False

    first_heading_idx = min(heading_indices.values())
    _fill_header_area(paragraphs, first_heading_idx, sections)

    ordered = sorted(heading_indices.items(), key=lambda item: item[1])
    for i, (section_key, heading_idx) in enumerate(ordered):
        end_idx = ordered[i + 1][1] if i + 1 < len(ordered) else len(paragraphs)
        lines = _section_lines(section_key, sections.get(section_key, ""))
        if lines:
            _write_section_between_headings(paragraphs, heading_idx, end_idx, lines, section_key)

    return True
