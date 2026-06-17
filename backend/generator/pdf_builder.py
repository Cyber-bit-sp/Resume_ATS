from io import BytesIO
from pathlib import Path


def _extract_lines_from_generated_docx(generation):
    if not generation.generated_file:
        return []

    suffix = Path(generation.generated_file.name).suffix.lower()
    if suffix != ".docx":
        return []

    try:
        from docx import Document

        with generation.generated_file.open("rb") as handle:
            doc = Document(handle)
    except Exception:
        return []

    lines = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = (paragraph.text or "").strip()
                    if text:
                        lines.append(text)

    return lines


def _extract_lines_from_generated_text(generation):
    lines = []
    for raw in (generation.generated_resume_text or "").splitlines():
        text = raw.strip()
        if text:
            lines.append(text)
    return lines


def build_resume_pdf_bytes(generation):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import simpleSplit
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("PDF download requires the reportlab package on the backend.") from exc

    lines = _extract_lines_from_generated_docx(generation) or _extract_lines_from_generated_text(generation)

    buffer = BytesIO()
    page_width, page_height = A4
    left = 50
    right = page_width - 50
    top = page_height - 50
    bottom = 50

    pdf = canvas.Canvas(buffer, pagesize=A4)
    y = top

    def write_line(text):
        nonlocal y
        is_heading = text.isupper() or text in {
            "Full Name",
            "Contact",
            "Professional Summary",
            "Technical Skills",
            "Professional Experience",
            "Projects",
            "Education",
            "Certifications",
        }
        font_name = "Helvetica-Bold" if is_heading else "Helvetica"
        font_size = 11 if is_heading else 10
        line_height = 15 if is_heading else 13
        wrapped = simpleSplit(text, font_name, font_size, right - left)
        for line in wrapped:
            if y < bottom:
                pdf.showPage()
                y = top
            pdf.setFont(font_name, font_size)
            pdf.drawString(left, y, line)
            y -= line_height

    for line in lines:
        write_line(line)

    pdf.save()
    buffer.seek(0)
    return buffer
